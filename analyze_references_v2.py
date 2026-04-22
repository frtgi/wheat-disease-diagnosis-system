#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
开题报告参考文献分析脚本 - v2
"""
import re
from docx import Document
from collections import defaultdict, Counter


def read_docx_content(file_path):
    """读取Word文档完整内容"""
    doc = Document(file_path)
    full_text = []
    paragraphs = []
    
    i = 0
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
            paragraphs.append({
                'index': i,
                'text': para.text,
                'style': para.style.name if para.style else ''
            })
            i += 1
    
    return '\n'.join(full_text), paragraphs


def extract_references(paragraphs):
    """提取所有参考文献信息"""
    references = []
    
    # 直接从原始段落索引中提取
    # 参考文献从原始段落索引88开始到108结束
    # 构建一个包含所有非空段落的列表，保留原始索引信息
    all_paragraphs = []
    doc = Document("/workspace/开题报告_2210122288 孙敬育_20260107.docx")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            all_paragraphs.append({
                'original_idx': i,
                'text': para.text.strip()
            })
    
    # 从原始索引88到108提取参考文献
    ref_start = 88
    ref_end = 108
    
    for orig_idx in range(ref_start, ref_end + 1):
        # 在all_paragraphs中找到对应原始索引的段落
        for para in all_paragraphs:
            if para['original_idx'] == orig_idx:
                ref_num = orig_idx - ref_start + 1
                references.append({
                    'number': ref_num,
                    'content': para['text'],
                    'raw_text': para['text']
                })
                break
    
    return references


def parse_reference_info(ref):
    """解析参考文献的详细信息"""
    info = {
        'number': ref['number'],
        'raw_text': ref['raw_text'],
        'content': ref['content'],
        'type': 'unknown',
        'authors': [],
        'title': '',
        'source': '',
        'year': '',
        'field': ''
    }
    
    text = ref['content']
    
    # 尝试识别参考文献类型
    if '[J]' in text:
        info['type'] = '期刊论文'
    elif '[D]' in text:
        info['type'] = '学位论文'
    elif '[P]' in text:
        info['type'] = '专利'
    elif '[C]' in text or '会议' in text:
        info['type'] = '会议论文'
    elif '[M]' in text:
        info['type'] = '图书'
    elif '[EB/OL]' in text or 'http' in text.lower():
        info['type'] = '网络资源'
    else:
        # 通过关键词判断
        if '期刊' in text or '学报' in text or '杂志' in text:
            info['type'] = '期刊论文'
        elif '学位论文' in text or '硕士' in text or '博士' in text or '大学' in text:
            info['type'] = '学位论文'
        elif '出版社' in text:
            info['type'] = '图书'
        else:
            info['type'] = '其他'
    
    # 提取年份
    year_match = re.search(r'(\d{4})', text)
    if year_match:
        info['year'] = year_match.group(1)
    
    # 识别领域
    if '小麦' in text or '病害' in text or '农业' in text or '植保' in text or '柑橘' in text or '病虫害' in text:
        info['field'] = '农业/植物保护'
    elif '深度学习' in text or '神经网络' in text or 'CNN' in text or 'YOLO' in text:
        info['field'] = '计算机视觉/深度学习'
    elif '多模态' in text or '融合' in text:
        info['field'] = '多模态融合'
    elif '知识图谱' in text or 'Graph' in text:
        info['field'] = '知识图谱'
    elif '诊断' in text or '检测' in text or '临床' in text or '银屑病' in text or '乳腺癌' in text:
        info['field'] = '智能诊断/检测'
    elif '决策' in text or '决策支持' in text:
        info['field'] = '决策支持系统'
    elif '生态' in text:
        info['field'] = '生态治理'
    elif '铁路' in text or '建筑' in text:
        info['field'] = '工程技术'
    elif '航空' in text:
        info['field'] = '航空航天'
    elif '企业' in text or '营销' in text:
        info['field'] = '企业管理'
    else:
        info['field'] = '其他'
    
    # 提取标题（尝试）
    title_match = re.search(r'[.，]\s*([^.，]+?)[.，]\s*(?:\[J\]|\[D\]|\[C\]|\[M\])', text)
    if title_match:
        info['title'] = title_match.group(1).strip()
    else:
        # 简单截取
        parts = text.split('.')
        if len(parts) > 1:
            info['title'] = parts[1].strip()
    
    return info


def find_citation_positions(paragraphs, references):
    """查找参考文献在正文中的引用位置"""
    citations = defaultdict(list)
    
    for para in paragraphs:
        text = para['text']
        para_index = para['index']
        
        # 查找引用标记 [数字]
        citation_matches = re.finditer(r'\[(\d+)\]', text)
        
        for match in citation_matches:
            ref_num = int(match.group(1))
            
            # 只关注1-21的参考文献
            if 1 <= ref_num <= 21:
                # 获取上下文（前后各100个字符）
                start = max(0, match.start() - 100)
                end = min(len(text), match.end() + 100)
                context = text[start:end]
                
                citations[ref_num].append({
                    'paragraph_index': para_index,
                    'paragraph_text': text,
                    'context': context,
                    'position': match.start()
                })
    
    return citations


def analyze_references(doc_path, output_path):
    """主分析函数"""
    print(f"正在读取文档: {doc_path}")
    full_text, paragraphs = read_docx_content(doc_path)
    
    print("提取参考文献...")
    references = extract_references(paragraphs)
    print(f"找到 {len(references)} 篇参考文献")
    
    print("解析参考文献信息...")
    parsed_refs = []
    for ref in references:
        parsed_ref = parse_reference_info(ref)
        parsed_refs.append(parsed_ref)
    
    print("查找引用位置...")
    citations = find_citation_positions(paragraphs, references)
    
    print("生成分析报告...")
    report = []
    report.append("=" * 80)
    report.append("开题报告参考文献详细分析报告")
    report.append("=" * 80)
    report.append(f"分析日期: 2026-04-22")
    report.append(f"参考文献总数: {len(references)} 篇")
    report.append("")
    
    # 1. 参考文献详细信息
    report.append("-" * 80)
    report.append("一、参考文献详细信息")
    report.append("-" * 80)
    for ref in parsed_refs:
        report.append(f"\n[{ref['number']}] {ref['content']}")
        report.append(f"   类型: {ref['type']}")
        report.append(f"   年份: {ref['year']}")
        report.append(f"   领域: {ref['field']}")
        if ref['title']:
            report.append(f"   标题: {ref['title']}")
    
    # 2. 引用位置和上下文
    report.append("\n" + "-" * 80)
    report.append("二、参考文献引用位置和上下文")
    report.append("-" * 80)
    for ref_num in sorted(citations.keys()):
        report.append(f"\n参考文献 [{ref_num}] 被引用 {len(citations[ref_num])} 次:")
        for i, cite in enumerate(citations[ref_num], 1):
            report.append(f"  引用位置 {i}:")
            report.append(f"    段落索引: {cite['paragraph_index']}")
            report.append(f"    上下文: {cite['context']}")
    
    # 3. 引用格式说明
    report.append("\n" + "-" * 80)
    report.append("三、引用格式说明")
    report.append("-" * 80)
    report.append("\n引用格式采用顺序编码制（GB/T 7714-2015）：")
    report.append("1. 正文引用格式: [序号]")
    report.append("2. 期刊论文格式: [序号] 作者. 标题[J]. 期刊名, 年份, 卷(期): 页码.")
    report.append("3. 学位论文格式: [序号] 作者. 标题[D]. 学校所在地: 学校名, 年份.")
    report.append("4. 外文文献格式: [序号] Author. Title[J]. Journal, Year, Volume(Issue): Pages.")
    report.append("5. 正文中的引用按出现顺序编号")
    
    # 4. 参考文献领域分布统计
    report.append("\n" + "-" * 80)
    report.append("四、参考文献领域分布统计")
    report.append("-" * 80)
    
    # 类型统计
    type_counts = Counter(ref['type'] for ref in parsed_refs)
    report.append("\n1. 参考文献类型分布:")
    for ref_type, count in sorted(type_counts.items()):
        report.append(f"   {ref_type}: {count} 篇 ({count/len(references)*100:.1f}%)")
    
    # 领域统计
    field_counts = Counter(ref['field'] for ref in parsed_refs)
    report.append("\n2. 参考文献领域分布:")
    for field, count in sorted(field_counts.items()):
        report.append(f"   {field}: {count} 篇 ({count/len(references)*100:.1f}%)")
    
    # 年份统计
    year_counts = Counter(ref['year'] for ref in parsed_refs if ref['year'])
    report.append("\n3. 参考文献年份分布:")
    for year in sorted(year_counts.keys()):
        report.append(f"   {year}年: {year_counts[year]} 篇")
    
    # 5. 参考文献与引用位置映射关系
    report.append("\n" + "-" * 80)
    report.append("五、参考文献与引用位置映射关系")
    report.append("-" * 80)
    
    for ref in parsed_refs:
        ref_num = ref['number']
        cite_count = len(citations.get(ref_num, []))
        title_display = ref['title'] if ref['title'] else ref['content'][:60]
        report.append(f"\n[{ref_num}] {title_display}...")
        report.append(f"   类型: {ref['type']} | 年份: {ref['year']} | 领域: {ref['field']}")
        report.append(f"   引用次数: {cite_count}")
        if citations.get(ref_num):
            report.append(f"   引用段落索引: {[c['paragraph_index'] for c in citations[ref_num]]}")
    
    # 6. 引用频次统计
    report.append("\n" + "-" * 80)
    report.append("六、参考文献引用频次排行")
    report.append("-" * 80)
    
    cite_freq = []
    for ref in parsed_refs:
        ref_num = ref['number']
        count = len(citations.get(ref_num, []))
        cite_freq.append((ref_num, count, ref['title'] if ref['title'] else ref['content'][:40]))
    
    cite_freq.sort(key=lambda x: x[1], reverse=True)
    for i, (ref_num, count, title) in enumerate(cite_freq, 1):
        status = "高频引用" if count >= 3 else "中频引用" if count >= 1 else "未引用"
        report.append(f"{i}. [{ref_num}] {title}... - {count}次 ({status})")
    
    # 保存报告
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(report))
    
    print(f"分析完成！结果已保存到: {output_path}")
    
    return parsed_refs, citations


if __name__ == "__main__":
    doc_path = "/workspace/开题报告_2210122288 孙敬育_20260107.docx"
    output_path = "/workspace/参考文献详细分析.txt"
    
    analyze_references(doc_path, output_path)
