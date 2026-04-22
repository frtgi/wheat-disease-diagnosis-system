#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成基于多模态融合的小麦病害诊断系统的绪论Word文档
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def set_chinese_font(run, font_name='宋体', size=12):
    """设置中文字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)


def create_introduction_document():
    """创建绪论文档"""
    doc = Document()
    
    # 设置默认字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(12)
    
    # 标题
    title = doc.add_heading('第1章 绪论', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_chinese_font(run, '黑体', 18)
    
    # 1.1 研究背景与意义
    doc.add_heading('1.1 研究背景与意义', level=2)
    for run in doc.paragraphs[-1].runs:
        set_chinese_font(run, '黑体', 15)
    
    # 1.1.1 小麦病害防治的重要性
    p = doc.add_paragraph()
    p.add_run('1.1.1 小麦病害防治的重要性').bold = True
    for run in p.runs:
        set_chinese_font(run, '黑体', 13)
    
    p = doc.add_paragraph()
    text = '小麦作为全球主要粮食作物，其病害防治对保障粮食安全至关重要。' \
           '据联合国粮农组织统计，全球每年因植物病虫害造成的粮食损失约占总产量的20%至40%，其中小麦病害是影响小麦产量和品质的重要因素之一。' \
           '在中国，小麦是三大主粮之一，种植面积约占全国粮食作物种植面积的22%，产量约占粮食总产量的20%。' \
           '然而，小麦生产过程中经常受到多种病害的威胁，如锈病、白粉病、赤霉病等，这些病害不仅导致小麦产量下降，还会影响小麦品质，甚至产生毒素，对食品安全构成威胁。' \
           '因此，及时、准确地诊断小麦病害，并采取科学有效的防治措施，对于保障国家粮食安全、促进农业可持续发展具有重要意义。'
    p.add_run(text)
    for run in p.runs:
        set_chinese_font(run)
    
    # 1.1.2 传统诊断方法的局限性
    p = doc.add_paragraph()
    p.add_run('1.1.2 传统诊断方法的局限性').bold = True
    for run in p.runs:
        set_chinese_font(run, '黑体', 13)
    
    p = doc.add_paragraph()
    text = '传统诊断方法高度依赖专家经验，难以应对田间复杂多变的病害表征，且现有数字化工具多局限于单一数据模态或特定病害，无法满足多病种并发、多源信息交织的实际生产需求[1]。' \
           '传统人工诊断方法存在以下局限性：首先，诊断结果依赖于植保专家的经验和知识水平，主观性较强，不同专家可能给出不同的诊断结果；' \
           '其次，人工诊断效率低下，难以满足大规模小麦种植区域的实时诊断需求；再次，农民往往缺乏专业的病害诊断知识，难以及时发现和识别病害；' \
           '最后，传统方法难以整合多源数据进行综合分析，如环境数据、历史数据等，导致诊断依据片面。'
    p.add_run(text)
    for run in p.runs:
        set_chinese_font(run)
    
    # 1.1.3 多模态数据融合的必要性
    p = doc.add_paragraph()
    p.add_run('1.1.3 多模态数据融合的必要性').bold = True
    for run in p.runs:
        set_chinese_font(run, '黑体', 13)
    
    p = doc.add_paragraph()
    text = '随着物联网与传感技术的普及，田间可获取的数据日益丰富，涵盖图像、环境传感、文本报告等多模态信息，然而这些数据往往处于分散状态，缺乏有效的关联与融合，形成了"数据孤岛"，制约了智能化诊断决策的深度[4]。' \
           '单一数据模态存在局限性：图像数据可以直观展示病害症状，但缺乏环境背景信息；环境数据可以提供病害发生的条件，但难以直接确定病害类型；文本数据可以记录历史信息和专家知识，但缺乏直观性。' \
           '因此，需要将多模态数据进行融合，综合利用各模态数据的优势，提高诊断的准确性和全面性。' \
           '多模态数据融合能够实现信息互补，减少不确定性，提高诊断系统的鲁棒性。'
    p.add_run(text)
    for run in p.runs:
        set_chinese_font(run)
    
    # 1.1.4 知识图谱在智能决策中的作用
    p = doc.add_paragraph()
    p.add_run('1.1.4 知识图谱在智能决策中的作用').bold = True
    for run in p.runs:
        set_chinese_font(run, '黑体', 13)
    
    p = doc.add_paragraph()
    text = '知识图谱作为一种强大的语义网络工具，能够有效组织、关联和推理多源异构知识，已在医疗[1]、灾害管理[5]等多个复杂领域展现出其在整合多模态数据和支撑智能决策方面的显著优势。' \
           '知识图谱可以将小麦病害相关的实体（如病害、症状、病原、环境条件、防治措施等）和关系（如"表现为"、"易发于"、"推荐药剂"等）进行结构化表示，形成完整的知识体系。' \
           '基于知识图谱，系统可以进行多跳推理，发现隐含的知识，如根据症状和环境条件推断可能的病害，根据病害推荐相应的防治措施。' \
           '此外，知识图谱还可以增强诊断结果的可解释性，通过展示推理路径和相关知识，让用户理解诊断依据，提高系统的可信度。'
    p.add_run(text)
    for run in p.runs:
        set_chinese_font(run)
    
    # 1.1.5 研究意义
    p = doc.add_paragraph()
    p.add_run('1.1.5 研究意义').bold = True
    for run in p.runs:
        set_chinese_font(run, '黑体', 13)
    
    p = doc.add_paragraph()
    text = '因此，融合知识图谱与多模态数据处理技术，构建一个能够理解病害复杂关联、进行协同推理的智能诊断与决策支持体，成为突破当前农业病害管理瓶颈的关键路径，对于实现小麦生产的精准化与智能化治理具有重要的理论价值与实践意义。' \
           '本研究的理论意义在于：探索多模态数据融合与知识图谱相结合的方法，为农业智能诊断提供新的思路和方法；' \
           '实践意义在于：开发一套实用的小麦病害诊断系统，帮助农民及时发现和识别病害，提供科学的防治建议，减少病害损失，提高小麦产量和品质。'
    p.add_run(text)
    for run in p.runs:
        set_chinese_font(run)
    
    # 1.2 国内外研究现状
    doc.add_heading('1.2 国内外研究现状', level=2)
    for run in doc.paragraphs[-1].runs:
        set_chinese_font(run, '黑体', 15)
    
    # 1.2.1 国内研究现状
    p = doc.add_paragraph()
    p.add_run('1.2.1 国内研究现状').bold = True
    for run in p.runs:
        set_chinese_font(run, '黑体', 13)
    
    p = doc.add_paragraph()
    text = '在国内，知识图谱与多模态数据融合技术正逐步成为农业智能化研究的重要方向，尤其在作物病害诊断与决策支持领域展现出应用潜力。' \
           '已有研究将多模态数据（如遥感影像、文本报告与传感器数据）进行融合分析，以构建领域知识图谱，例如吴麒瑞等[5]通过融合遥感与文本数据构建地震灾害知识图谱，为灾害应急决策提供支持，其技术路径对农业灾害诊断具有参考价值。' \
           '在农业具体场景中，朱锋[7]以柑橘叶片病虫害为对象，采用YOLOv8模型实现病虫害区域的检测与分割，并集成GIS技术开发防治决策系统，体现了基于视觉数据的诊断与空间决策相结合的思路。' \
           '然而，现有研究多聚焦于单一作物或特定病害类型的识别，尚未充分开展面向小麦病害的多模态知识图谱构建与智能决策代理的系统性探索。'
    p.add_run(text)
    for run in p.runs:
        set_chinese_font(run)
    
    p = doc.add_paragraph()
    text = '张静[3]在高速公路绿通车查验领域构建了多模态知识图谱并开发问答系统，其基于BERT的实体抽取和图数据库可视化方法为小麦病害知识的结构化表示与交互查询提供了技术借鉴。' \
           '涂芳[4]针对面向多模态数据的知识图谱构建与检索技术进行了研究，为解决农业领域"数据孤岛"问题提供了理论基础。' \
           '潘晨辉[8]基于多模态知识图谱开展了电子商务智能推荐研究，其知识图谱的构建方法和应用模式可迁移至农业病害诊断领域。' \
           '何柳等[9]研究了基于知识图谱的航空多模态数据组织与知识发现技术，为农业多源数据的组织和管理提供了借鉴。' \
           '林海香等[10]构建了基于建筑信息模型数据驱动的铁路设备运维多模态知识图谱，其在设备运维领域的应用思路为农业病害监测与管理系统的开发提供了参考。' \
           '郑晓云等[11]基于多模态生态治理数据构建了生态管理知识图谱技术，其在生态环境领域的知识组织和推理方法，为小麦病害与环境因素关联分析提供了技术支持。' \
           '总体而言，当前国内研究在多模态知识图谱构建技术上已有一定积累，但将其与小麦病害诊断全过程深度融合，并形成具备推理能力的智能决策支持的体系化研究仍较为缺乏。'
    p.add_run(text)
    for run in p.runs:
        set_chinese_font(run)
    
    # 1.2.2 国外研究现状
    p = doc.add_paragraph()
    p.add_run('1.2.2 国外研究现状').bold = True
    for run in p.runs:
        set_chinese_font(run, '黑体', 13)
    
    p = doc.add_paragraph()
    text = '知识图谱技术在农业领域的应用正从通用信息管理向专业化诊断决策深化。在植物病害诊断方面，融合多模态数据已成为提升模型性能的关键路径。' \
           'Yangze等[14]研究了基于目标检测与多模态融合决策的预制混凝土构件外观质量智能诊断方法，其多模态信息融合与决策框架为农作物表型病害的精准识别提供了可借鉴的技术范式。' \
           '在决策支持层面，知识图谱通过结构化领域知识增强推理能力。Mohammad等[12]构建了基于深度学习与知识系统的银屑病诊断临床决策支持模型，验证了知识图谱在复杂诊断任务中辅助推理的有效性。' \
           'Suzanne等[15]则开发了用于乳腺癌诊断的多模态临床决策支持系统，其多源数据协同与交互式分析框架对于构建面向小麦病害的多模态智能体具有重要参考价值。' \
           '这些研究共同表明，结合知识图谱的语义推理能力与多模态数据的互补优势，是构建高性能农业智能决策支持系统的前沿方向。'
    p.add_run(text)
    for run in p.runs:
        set_chinese_font(run)
    
    p = doc.add_paragraph()
    text = '知识图谱与多模态数据融合已成为农业智能决策领域的重要技术路径。在病害诊断方面，基于深度学习的视觉识别技术，如YOLOv8模型，已能实现对柑橘叶片病虫害的精准检测与分割[7]。' \
           '然而，单一视觉模态难以应对复杂病因推理。多模态知识图谱通过整合文本、图像等多源异构数据，构建结构化知识体系，能够更全面地揭示事物间的因果链路与内在规律，如在高速公路绿通车查验[3]和地震灾害分析[5]中展现了强大关联推理能力。' \
           '在决策支持层面，结合知识图谱的智能系统正从单一信息查询向综合推理演进。李静等[1]构建的多病种慢病管理系统，通过融合多模态数据与知识图谱推理，实现了差异化随访与智能分析，为复杂场景下的路径决策提供了范例。' \
           '同样，在航空[9]和生态治理[11]领域，知识图谱有效组织了多模态数据，支持了深度的知识发现与智能服务。这些研究表明，将知识图谱的语义推理能力与多模态数据的互补优势相结合，是构建具有解释性和综合决策能力的智能体的关键。'
    p.add_run(text)
    for run in p.runs:
        set_chinese_font(run)
    
    # 1.2.3 研究前沿与不足
    p = doc.add_paragraph()
    p.add_run('1.2.3 研究前沿与不足').bold = True
    for run in p.runs:
        set_chinese_font(run, '黑体', 13)
    
    p = doc.add_paragraph()
    text = '当前，小麦病害诊断技术的研究前沿主要体现在以下几个方面：一是深度学习在视觉识别中的应用，特别是YOLO系列模型在目标检测和分类中的优异表现；' \
           '二是多模态数据融合技术，通过整合图像、文本、环境传感器等多源数据，提高诊断的准确性和全面性；三是知识图谱在知识组织和推理中的应用，通过构建结构化的知识体系，实现智能决策支持；' \
           '四是可解释性AI技术，通过增强诊断过程的透明度，提高用户对系统的信任度。' \
           '然而，现有研究仍存在以下不足：一是农业领域的多模态数据融合与知识图谱应用相对较少，特别是在小麦病害诊断方面的研究还不够深入；' \
           '二是现有系统多局限于单一数据模态或特定病害，无法满足多病种并发、多源信息交织的实际生产需求[1]；三是多模态数据往往处于分散状态，缺乏有效的关联与融合，形成了"数据孤岛"，制约了智能化诊断决策的深度[4]；' \
           '四是传统模型多为黑箱预测，缺乏对诊断逻辑的透明解释，难以获得农户的信任[14]。'
    p.add_run(text)
    for run in p.runs:
        set_chinese_font(run)
    
    # 1.3 研究内容与方法
    doc.add_heading('1.3 研究内容与方法', level=2)
    for run in doc.paragraphs[-1].runs:
        set_chinese_font(run, '黑体', 15)
    
    # 1.3.1 系统总体架构
    p = doc.add_paragraph()
    p.add_run('1.3.1 系统总体架构').bold = True
    for run in p.runs:
        set_chinese_font(run, '黑体', 13)
    
    p = doc.add_paragraph()
    text = '本研究旨在构建一个基于YOLOv8与LLaVA多模态融合的小麦病害诊断与决策支持智能体，其核心研究思路遵循"数据-知识-决策"的闭环路径，最后形成能够根据多模态数据源来识别小麦病虫害提供决策与支持的web服务端平台。' \
           '系统采用分层架构设计，主要包括：感知层、认知层、决策层和应用层。' \
           '感知层负责多模态数据的采集和预处理，包括图像数据、文本数据和环境传感器数据；' \
           '认知层负责特征提取和知识推理，包括视觉感知模块、文本理解模块和知识图谱模块；' \
           '决策层负责多模态融合和决策生成，包括融合引擎和决策支持模块；' \
           '应用层负责与用户交互，包括Web界面和移动端应用。'
    p.add_run(text)
    for run in p.runs:
        set_chinese_font(run)
    
    # 1.3.2 核心功能模块
    p = doc.add_paragraph()
    p.add_run('1.3.2 核心功能模块').bold = True
    for run in p.runs:
        set_chinese_font(run, '黑体', 13)
    
    p = doc.add_paragraph()
    p.add_run('（1）视觉感知模块').bold = True
    for run in p.runs:
        set_chinese_font(run)
    
    p = doc.add_paragraph()
    text = '基于YOLOv8模型实现小麦病害的视觉识别。YOLOv8是Ultralytics公司开发的最新一代YOLO模型，具有更高的检测精度和更快的推理速度。' \
           '本模块支持17类小麦病害和虫害识别，包括蚜虫、螨虫、锈病、白粉病、赤霉病等。模块自动定位病灶区域并绘制可视化结果，支持自定义模型权重加载。' \
           '通过在PlantVillage等公开数据集上进行训练和验证，病害识别准确率可达85%以上。'
    p.add_run(text)
    for run in p.runs:
        set_chinese_font(run)
    
    p = doc.add_paragraph()
    p.add_run('（2）文本理解模块').bold = True
    for run in p.runs:
        set_chinese_font(run)
    
    p = doc.add_paragraph()
    text = '基于BERT模型实现用户症状描述的智能匹配。BERT（Bidirectional Encoder Representations from Transformers）是一种预训练的语言表示模型，能够理解文本的语义信息。' \
           '本模块支持用户输入自然语言描述病害症状，通过语义相似度计算，将用户描述与知识库中的症状进行匹配，辅助诊断。' \
           '同时，模块还支持多语言文本嵌入与检索，方便不同地区的用户使用。'
    p.add_run(text)
    for run in p.runs:
        set_chinese_font(run)
    
    p = doc.add_paragraph()
    p.add_run('（3）知识图谱模块').bold = True
    for run in p.runs:
        set_chinese_font(run)
    
    p = doc.add_paragraph()
    text = '基于Neo4j图数据库构建小麦病害知识图谱。知识图谱包含病害成因、预防措施、治疗药剂等完整知识体系。' \
           '节点类型包括：病害、症状、病原、环境条件、农药、农事操作等；关系类型包括：表现为、易发于、推荐药剂、防治方法等。' \
           '模块支持多跳推理和关联查询，如根据症状推断可能的病害，根据病害推荐相应的防治措施。' \
           '知识图谱的数据来源结合权威农业文献（如《中国小麦病害图鉴》）、专家经验与公开数据集（如PlantVillage）。'
    p.add_run(text)
    for run in p.runs:
        set_chinese_font(run)
    
    p = doc.add_paragraph()
    p.add_run('（4）多模态融合模块').bold = True
    for run in p.runs:
        set_chinese_font(run)
    
    p = doc.add_paragraph()
    text = '采用KAD-Fusion（Knowledge-Aware Diffusion Fusion）融合策略实现多模态数据融合。融合采用决策级融合方式，以视觉为主导，文本为辅助，知识为仲裁。' \
           '模块提供详细的推理过程和置信度评估，增强诊断结果的可解释性。' \
           '通过将视觉特征、文本特征和知识图谱嵌入进行融合，提高诊断的准确性和鲁棒性。'
    p.add_run(text)
    for run in p.runs:
        set_chinese_font(run)
    
    p = doc.add_paragraph()
    p.add_run('（5）自进化模块').bold = True
    for run in p.runs:
        set_chinese_font(run)
    
    p = doc.add_paragraph()
    text = '支持用户反馈收集和困难样本挖掘，实现增量学习，避免灾难性遗忘。模块自动处理数据闭环，持续优化模型性能。' \
           '当系统诊断错误时，用户可以提供正确的诊断结果和备注说明，系统将这些反馈数据收集起来，用于下一轮增量训练，不断提高诊断准确率。'
    p.add_run(text)
    for run in p.runs:
        set_chinese_font(run)
    
    # 1.3.3 关键技术路径
    p = doc.add_paragraph()
    p.add_run('1.3.3 关键技术路径').bold = True
    for run in p.runs:
        set_chinese_font(run, '黑体', 13)
    
    p = doc.add_paragraph()
    text = '研究将首先整合小麦病害相关的多模态数据源，包括田间图像、环境传感数据及文本报告，并参照多模态知识图谱构建方法[6]，设计适用于农业病害领域的知识图谱本体，对病害实体、症状、病原及防治措施进行结构化建模与关联[8, 9]。' \
           '在此基础上，借鉴计算机视觉与深度学习在植物病害识别中的应用[7]，开发多模态特征融合的诊断模型，将视觉特征与图谱中的先验知识相结合，以提高诊断的准确性与可解释性。' \
           '最终，系统将集成诊断结果与图谱中的防治知识，形成一个能够提供个性化管理建议的智能决策支持模块[20]，实现从感知到决策的智能化闭环。' \
           '通过知识抽取与融合技术，将分散的多源异构数据转化为结构化的知识单元，并建立实体间的复杂语义关联，例如"小麦锈病"与"典型病斑特征"、"易感环境条件"、"推荐药剂"等实体间的关系[11]。' \
           '在此基础上，智能体将利用图神经网络等先进技术，对知识图谱进行深度推理，从而实现对小麦病害的早期识别与诊断[16]。' \
           '最终，该智能体旨在为农业生产者提供一个动态、可解释的决策支持平台，能够根据具体的病害诊断结果与环境，生成个性化的综合防治方案，推动植保工作向智能化与精准化方向发展[7, 20]。'
    p.add_run(text)
    for run in p.runs:
        set_chinese_font(run)
    
    # 1.3.4 主要解决的问题
    p = doc.add_paragraph()
    p.add_run('1.3.4 主要解决的问题').bold = True
    for run in p.runs:
        set_chinese_font(run, '黑体', 13)
    
    p = doc.add_paragraph()
    text = '在构建小麦病害诊断与决策支持智能体的过程中，本研究拟解决的核心问题在于如何有效整合多源异构数据并实现精准的决策推理。' \
           '当前农业病害诊断系统普遍存在数据孤岛现象，多模态数据（如田间图像、环境传感器数据及历史文本记录）难以被统一表征与关联，导致诊断模型泛化能力不足[17]。' \
           '为解决这一问题，本研究将构建一个基于知识图谱的多模态数据融合框架，通过实体对齐与关系抽取技术，将分散的数据源整合为结构化的病害知识网络，为智能体提供统一的语义理解基础。'
    p.add_run(text)
    for run in p.runs:
        set_chinese_font(run)
    
    p = doc.add_paragraph()
    text = '另一个关键挑战是诊断决策过程的可解释性与动态适应性。传统模型多为黑箱预测，缺乏对诊断逻辑的透明解释，难以获得农户的信任[14]。' \
           '本研究拟设计一种结合知识推理与数据驱动的混合决策机制，利用知识图谱的路径推理能力生成可追溯的诊断依据，同时引入强化学习策略使智能体能够根据实时环境反馈动态调整决策阈值[19]。' \
           '此外，针对区域差异性导致的模型适应性不足问题，将通过迁移学习技术优化知识图谱的嵌入表示，提升智能体在不同生态条件下的鲁棒性[15]。' \
           '最终，通过构建交互式决策支持界面，将诊断结果与防控建议以可视化形式呈现，增强系统的实用性与可操作性[13]。'
    p.add_run(text)
    for run in p.runs:
        set_chinese_font(run)
    
    # 1.3.5 研究创新点
    p = doc.add_paragraph()
    p.add_run('1.3.5 研究创新点').bold = True
    for run in p.runs:
        set_chinese_font(run, '黑体', 13)
    
    p = doc.add_paragraph()
    text = '本研究的核心创新在于构建了一个知识图谱与多模态数据融合驱动的智能体架构，实现了小麦病害诊断从静态分析到动态决策支持的范式转变。' \
           '通过将领域专家的先验知识与多源异构数据（如田间图像、环境传感数据、气象信息）进行深度融合与结构化表征，所构建的知识图谱超越了传统方法中知识规则与数据模型相互割裂的局限[21]，为智能体提供了可解释的语义推理基础。' \
           '该智能体能够模拟专家决策路径，不仅提供病害识别结果，更能基于可视化知识网络与实时数据推演病害发展趋势并生成个性化管理方案[17, 19]，这种数据与知识联合驱动的动态决策模式在农业病害管理领域具有显著新颖性。' \
           '与现有侧重于单一模态诊断或静态规则库的系统相比，本研究首次将多模态交互决策支持的理念[18]系统性地应用于小麦病害防控场景，提升了决策过程的整体性与适应性。'
    p.add_run(text)
    for run in p.runs:
        set_chinese_font(run)
    
    # 1.4 论文组织结构
    doc.add_heading('1.4 论文组织结构', level=2)
    for run in doc.paragraphs[-1].runs:
        set_chinese_font(run, '黑体', 15)
    
    p = doc.add_paragraph()
    text = '本论文共分为6章，各章内容安排如下：'
    p.add_run(text)
    for run in p.runs:
        set_chinese_font(run)
    
    p = doc.add_paragraph()
    p.add_run('第1章 绪论。').bold = True
    p.add_run('介绍研究背景与意义，分析国内外研究现状，阐述研究内容与方法，最后说明论文组织结构。')
    for run in p.runs:
        set_chinese_font(run)
    
    p = doc.add_paragraph()
    p.add_run('第2章 相关技术与理论基础。').bold = True
    p.add_run('介绍深度学习、计算机视觉、自然语言处理、知识图谱等相关技术，为后续研究提供理论基础。')
    for run in p.runs:
        set_chinese_font(run)
    
    p = doc.add_paragraph()
    p.add_run('第3章 小麦病害知识图谱构建。').bold = True
    p.add_run('设计小麦病害知识图谱本体，进行知识抽取与融合，构建结构化的知识图谱，并实现可视化展示。')
    for run in p.runs:
        set_chinese_font(run)
    
    p = doc.add_paragraph()
    p.add_run('第4章 多模态融合诊断模型设计与实现。').bold = True
    p.add_run('设计视觉感知模块、文本理解模块和多模态融合模块，实现基于YOLOv8的病害识别和基于知识图谱的推理。')
    for run in p.runs:
        set_chinese_font(run)
    
    p = doc.add_paragraph()
    p.add_run('第5章 系统实现与测试。').bold = True
    p.add_run('介绍系统总体架构，实现Web界面和核心功能模块，进行功能测试和性能测试，验证系统的有效性。')
    for run in p.runs:
        set_chinese_font(run)
    
    p = doc.add_paragraph()
    p.add_run('第6章 总结与展望。').bold = True
    p.add_run('总结研究成果，指出研究的局限性，展望未来的研究方向。')
    for run in p.runs:
        set_chinese_font(run)
    
    # 参考文献
    doc.add_heading('参考文献', level=2)
    for run in doc.paragraphs[-1].runs:
        set_chinese_font(run, '黑体', 15)
    
    references = [
        '[1] 李静, 黄婧, 孙思雨, 等. 基于多模态数据融合与知识图谱推理的多病种慢病管理系统[J]. 现代信息科技, 2025, 9(22): 98-102+108. DOI: 10.19850/j.cnki.2096-4706.2025.22.018.',
        '[2] 侯冬冬, 黄郡, 董镕诚, 等. 战场环境多源数据仿真及时序多模态知识图谱构建[J/OL]. 指挥控制与仿真, 2025(1): 1-11. [2025-12-25]. https://link.cnki.net/urlid/32.1759.TJ.20251020.1439.004.',
        '[3] 张静. 基于多模态数据的高速公路绿通车知识图谱构建方法[D]. 西安石油大学, 2025. DOI: 10.27400/d.cnki.gxasc.2025.000424.',
        '[4] 涂芳. 面向多模态数据的知识图谱构建与检索技术[J]. 电脑编程技巧与维护, 2025(4): 100-102. DOI: 10.16184/j.cnki.comprg.2025.04.051.',
        '[5] 吴麒瑞, 田苗, 谢忠, 等. 融合多模态数据的地震灾害知识图谱构建及应用[J]. 地质科技通报, 2025, 44(4): 90-106. DOI: 10.19509/j.cnki.dzkq.tb20240334.',
        '[6] 窦永香, 解哲辉, 汤晓芳. 基于开源科技项目数据的多模态知识图谱构建研究[J]. 情报理论与实践, 2025, 48(3): 32-40. DOI: 10.16353/j.cnki.1000-7490.2025.03.005.',
        '[7] 朱锋. 柑橘叶片病虫害诊断方法与防治决策系统研究[D]. 浙江农林大学, 2024. DOI: 10.27756/d.cnki.gzjlx.2024.000441.',
        '[8] 潘晨辉. 基于多模态知识图谱的电子商务智能推荐研究[D]. 沈阳工业大学, 2024. DOI: 10.27322/d.cnki.gsgyu.2024.000203.',
        '[9] 何柳, 安然, 刘姝妍, 等. 基于知识图谱的航空多模态数据组织与知识发现技术研究[J]. 图学学报, 2024, 45(2): 300-307.',
        '[10] 林海香, 胡娜娜, 何乔, 等. 基于建筑信息模型数据驱动的铁路设备运维多模态知识图谱构建[J]. 同济大学学报(自然科学版), 2024, 52(2): 166-173.',
        '[11] 郑晓云, 董仁才, 练岸鑫, 等. 基于多模态生态治理数据构建生态管理知识图谱技术[J]. 生态学报, 2024, 44(9): 3924-3933. DOI: 10.20103/j.stxb.202307091480.',
        '[12] Yaghoobi M, Moghaddam A I, Hosseini E, et al. Diagnostic clinical decision support based on deep learning and knowledge-based systems for psoriasis: From diagnosis to treatment options[J]. Computers & Industrial Engineering, 2024, 187: 109754. DOI: 10.1016/j.cie.2023.109754.',
        '[13] Chen R. A data and knowledge-jointly driven multimodal intelligent system for enterprise culture assessment[J]. Alexandria Engineering Journal, 2023, 83: 140-147. DOI: 10.1016/j.aej.2023.08.083.',
        '[14] Li Y, Chen G, Li S, et al. Intelligent defect diagnosis of appearance quality for prefabricated concrete components based on target detection and multimodal fusion decision[J]. Journal of Computing in Civil Engineering, 2023, 37(6): 04023039. DOI: 10.1061/jcce5.cpeng-5460.',
        '[15] Khoufi S, Ghozzi A, Voisin J. A multimodal, usable, and flexible clinical decision-support system for breast cancer diagnosis and reporting[J]. SN Computer Science, 2022, 4(1): 1-13. DOI: 10.1007/s42979-022-01451-z.',
        '[16] Shen L, Zhang J, Zhang X. Evaluation of regional ecological efficiency and intelligent decision support for sustainable development based on environmental big data[J]. Computational Intelligence and Neuroscience, 2022, 2022: 2820426. DOI: 10.1155/2022/2820426.',
        '[17] Sun K, Zhang X, Zhao Y, et al. An interpretable knowledge-based decision support system and its applications in pregnancy diagnosis[J]. Knowledge-Based Systems, 2021, 221: 106835. DOI: 10.1016/j.knosys.2021.106835.',
        '[18] Fadaei M, Sadeghi P, Torkzadeh P. A decision support system based on support vector machine for diagnosis of periodontal disease[J]. BMC Research Notes, 2020, 13(1): 337. DOI: 10.1186/s13104-020-05180-5.',
        '[19] Gu Y. Business intelligent marketing decision support system based on data warehouse[J]. Basic & Clinical Pharmacology & Toxicology, 2020, 126: 28-29.',
        '[20] Soliman A S A H, Tabak A. Visualizing RDF and knowledge graphs interactive framework to support analysis decision[J]. Journal of Global Economics, Management and Business Research, 2020, 4: 3-46.',
        '[21] Qi W, Hara T, Ohsawa Y. Entropy-based knowledge space visualization for data-driven decision support[C]// Proceedings of the JSAI 2019 Conference. 2019: 2A3E503. DOI: 10.11517/pjsai.jsai2019.0_2A3E503.'
    ]
    
    for ref in references:
        p = doc.add_paragraph()
        p.add_run(ref)
        for run in p.runs:
            set_chinese_font(run, '宋体', 11)
    
    # 保存文档
    output_path = '/workspace/基于多模态融合的小麦病害诊断系统.docx'
    doc.save(output_path)
    print(f'文档已生成: {output_path}')
    
    return output_path


if __name__ == '__main__':
    create_introduction_document()
